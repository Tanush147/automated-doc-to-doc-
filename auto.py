import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import shutil
import subprocess
import tempfile
import platform
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Define paths
PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
CSV_PATH = os.path.join(PROJECT_ROOT, "Untitled Spreadsheet_Sheet1.csv")
TEMPLATE_PATH = os.path.join(PROJECT_ROOT, "Certificate.docx")
OUTPUT_FOLDER = os.path.join(PROJECT_ROOT, "certificates")
PDF_OUTPUT_FOLDER = os.path.join(PROJECT_ROOT, "certificatesPDF")
SYSTEM = platform.system()  # 'Darwin' (macOS), 'Windows', 'Linux'
KEEP_DOCX = True  # Keep generated DOCX to verify design and positioning
WORD_ONLY_PDF = True  # Enforce Microsoft Word-based PDF conversion for layout fidelity

# Create output folders
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(PDF_OUTPUT_FOLDER, exist_ok=True)

# Load CSV data
df = pd.read_csv(CSV_PATH)

# Initialize certificate counter
certificate_counter = 201

def replace_text_preserve_formatting(doc, old_text, new_text):
    """Replace only the placeholder text within existing runs to preserve layout.

    This avoids replacing entire XML blocks, which can shift positioning. If a
    placeholder spans multiple runs, we skip it and log a warning to prevent
    layout changes. Ensure placeholders in the template are contained within a
    single run for exact in-place replacement.
    """
    def replace_across_runs(paragraph, placeholder, replacement):
        # Concatenate all run texts
        runs = paragraph.runs
        concat = ''.join(r.text or '' for r in runs)
        start = concat.find(placeholder)
        if start == -1:
            return False
        end = start + len(placeholder)
        # Map run boundaries in concatenated text
        boundaries = []
        pos = 0
        for r in runs:
            txt = r.text or ''
            boundaries.append((pos, pos + len(txt)))
            pos += len(txt)
        # Apply replacement across overlapping runs
        first_done = False
        for i, r in enumerate(runs):
            rstart, rend = boundaries[i]
            if rend <= start or rstart >= end:
                # Run outside placeholder region; leave as is
                continue
            # Overlap exists
            local_start = max(0, start - rstart)
            local_end = min(len(r.text or ''), end - rstart)
            before = (r.text or '')[:local_start]
            after = (r.text or '')[local_end:]
            if not first_done:
                # Put full replacement into the first overlapping run
                r.text = before + replacement + after
                first_done = True
            else:
                # Remove placeholder fragment from subsequent runs
                r.text = before + after
        return True

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            found_in_run = False
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    found_in_run = True
            if not found_in_run:
                # Try cross-run replacement without altering layout
                replaced = replace_across_runs(paragraph, old_text, new_text)
                if not replaced:
                    print(
                        f"⚠️ Could not replace split placeholder '{old_text}' in a paragraph without layout changes."
                    )

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        found_in_run = False
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                found_in_run = True
                        if not found_in_run:
                            replaced = replace_across_runs(paragraph, old_text, new_text)
                            if not replaced:
                                print(
                                    f"⚠️ Could not replace split placeholder '{old_text}' in a table cell without layout changes."
                                )

# Process each row in the CSV
for _, row in df.iterrows():
    name = str(row['Name'])
    course_name = "Data Science"  # Static value for course_name
    count = str(certificate_counter)
    certificate_counter += 1  # Increment counter for next certificate

    # Load the template document
    doc = Document(TEMPLATE_PATH)

    # Replace placeholders using formatting-preserving method
    replace_text_preserve_formatting(doc, '{{name}}', name)
    replace_text_preserve_formatting(doc, '{{ name }}', name)
    replace_text_preserve_formatting(doc, '{{count}}', count)
    replace_text_preserve_formatting(doc, '{{ count }}', count)
    replace_text_preserve_formatting(doc, '{{course_name}}', course_name)
    replace_text_preserve_formatting(doc, '{{ course_name }}', course_name)

    # Save the personalized document in the certificates folder
    filename = f"Certificate_{count}_{name.replace(' ', '_')}.docx"
    output_path = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")

# Convert DOCX to PDF (robust, cross-platform flow)
conversion_ok = False

# Use a temporary local directory to avoid iCloud/quartz permission issues on macOS
with tempfile.TemporaryDirectory() as tmpdir:
    tmp_in = os.path.join(tmpdir, "in")
    tmp_out = os.path.join(tmpdir, "out")
    os.makedirs(tmp_in, exist_ok=True)
    os.makedirs(tmp_out, exist_ok=True)

    # Copy DOCX files to temp directory
    docx_files = [f for f in os.listdir(OUTPUT_FOLDER) if f.lower().endswith('.docx')]
    for f in docx_files:
        shutil.copy2(os.path.join(OUTPUT_FOLDER, f), os.path.join(tmp_in, f))

    # Try docx2pdf directory conversion first (fastest)
    if SYSTEM in ("Darwin", "Windows"):
        try:
            # Convert all files in the directory
            convert(tmp_in, tmp_out)
            # Move PDFs to final destination
            for pf in os.listdir(tmp_out):
                if pf.lower().endswith('.pdf'):
                    shutil.move(os.path.join(tmp_out, pf), os.path.join(PDF_OUTPUT_FOLDER, pf))
            # Check if conversion was successful
            produced = [f for f in os.listdir(PDF_OUTPUT_FOLDER) if f.lower().endswith('.pdf')]
            conversion_ok = len(produced) >= len(docx_files) and len(docx_files) > 0
        except Exception as e:
            print(f"docx2pdf directory conversion failed: {e}")
            conversion_ok = False

# Fallback: Try docx2pdf per-file conversion (still uses Word)
if not conversion_ok and SYSTEM in ("Darwin", "Windows"):
    try:
        converted_count = 0
        for f in os.listdir(OUTPUT_FOLDER):
            if f.lower().endswith('.docx'):
                src = os.path.join(OUTPUT_FOLDER, f)
                pdf_name = os.path.splitext(f)[0] + ".pdf"
                dst = os.path.join(PDF_OUTPUT_FOLDER, pdf_name)
                convert(src, dst)
                if os.path.exists(dst):
                    converted_count += 1
        conversion_ok = converted_count > 0
    except Exception as e:
        print(f"docx2pdf per-file conversion failed: {e}")
        conversion_ok = False

# Final fallback: Use LibreOffice unless Word-only conversion is enforced
if not conversion_ok and not WORD_ONLY_PDF:
    soffice = shutil.which('soffice')
    if soffice:
        print("⚠️ Using LibreOffice for PDF conversion. Note: layout may shift compared to Word-based conversion.")
        try:
            subprocess.run([
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", PDF_OUTPUT_FOLDER,
                *[os.path.join(OUTPUT_FOLDER, f) for f in os.listdir(OUTPUT_FOLDER) if f.lower().endswith('.docx')]
            ], check=True)
            produced = [f for f in os.listdir(PDF_OUTPUT_FOLDER) if f.lower().endswith('.pdf')]
            conversion_ok = len(produced) > 0
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")
            conversion_ok = False

if not conversion_ok and WORD_ONLY_PDF:
    print("❌ Word-only PDF conversion enforced. Install Microsoft Word and fonts used in the template to preserve layout and design.")

# Clean up and report results
if conversion_ok:
    if not KEEP_DOCX:
        try:
            shutil.rmtree(OUTPUT_FOLDER, ignore_errors=True)
        except Exception:
            pass
        print("✅ PDFs created in 'certificatesPDF'. 'certificates' folder removed.")
    else:
        print("✅ PDFs created in 'certificatesPDF'. DOCX kept in 'certificates'.")
else:
    if SYSTEM == "Darwin":
        hint = "On macOS, try running from a local non-iCloud folder or ensure LibreOffice is installed."
    elif SYSTEM == "Windows":
        hint = "On Windows, ensure Microsoft Word or LibreOffice is installed."
    else:
        hint = "On Linux, install LibreOffice (soffice) for conversion."
    print(f"❌ PDF conversion did not complete. {hint}")

print("✅ Process completed. End result: PDFs in 'certificatesPDF'.")
