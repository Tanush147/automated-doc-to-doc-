import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import shutil
import subprocess
import tempfile
import platform

# Resolve absolute paths (avoids iCloud/space path issues)
PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
CSV_PATH = os.path.join(PROJECT_ROOT, "Untitled Spreadsheet_Sheet1.csv")
TEMPLATE_PATH = os.path.join(PROJECT_ROOT, "INTERNSHIP OFFER LETTER 3.docx")
OUTPUT_FOLDER = os.path.join(PROJECT_ROOT, "offer_letters")
PDF_OUTPUT_FOLDER = os.path.join(PROJECT_ROOT, "offer_lettersPDF")
SYSTEM = platform.system()  # 'Darwin' (macOS), 'Windows', 'Linux'

# Ensure output folders exist
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(PDF_OUTPUT_FOLDER, exist_ok=True)

# Load data from the CSV file
df = pd.read_csv(CSV_PATH)

# Generate all DOCX first
for _, row in df.iterrows():
    title = str(row['Title'])
    name = str(row['Name'])

    # Open the template for each recipient
    doc = Document(TEMPLATE_PATH)

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        if '<<Title>>' in para.text or '<<Name>>' in para.text:
            para.text = para.text.replace('<<Title>>', title).replace('<<Name>>', name)

    # Replace placeholders in tables (if any)
    for table in doc.tables:
        for trow in table.rows:
            for cell in trow.cells:
                if '<<Title>>' in cell.text or '<<Name>>' in cell.text:
                    cell.text = cell.text.replace('<<Title>>', title).replace('<<Name>>', name)

    # Save the personalized document in the offer_letters folder
    filename = f"Offer_Letter_{name.replace(' ', '_')}.docx"
    output_path = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")

# Convert DOCX to PDF (robust, cross-platform flow)
conversion_ok = False

# Use a temporary local directory to avoid iCloud/quartz permission issues on macOS
with tempfile.TemporaryDirectory() as tmpdir:
    tmp_in = os.path.join(tmpdir, "in")
    tmp_out = os.path.join(tmpdir, "out")
    os.makedirs(tmp_in, exist_ok=True)
    os.makedirs(tmp_out, exist_ok=True)

    # Copy docx files into tmp_in
    docx_files = [f for f in os.listdir(OUTPUT_FOLDER) if f.lower().endswith('.docx')]
    for f in docx_files:
        shutil.copy2(os.path.join(OUTPUT_FOLDER, f), os.path.join(tmp_in, f))

    # Prefer docx2pdf on macOS/Windows; skip on Linux (no MS Word/Quartz)
    if SYSTEM in ("Darwin", "Windows"):
        try:
            # Directory conversion is more reliable
            convert(tmp_in, tmp_out)
            # Move PDFs back to the project PDF folder
            for pf in os.listdir(tmp_out):
                if pf.lower().endswith('.pdf'):
                    shutil.move(os.path.join(tmp_out, pf), os.path.join(PDF_OUTPUT_FOLDER, pf))
            # Verify conversion produced expected count
            produced = [f for f in os.listdir(PDF_OUTPUT_FOLDER) if f.lower().endswith('.pdf')]
            conversion_ok = len(produced) >= len(docx_files) and len(docx_files) > 0
        except Exception as e:
            print(f"docx2pdf directory conversion failed: {e}")
            conversion_ok = False

# Fallback: per-file conversion using docx2pdf (macOS/Windows)
if not conversion_ok and SYSTEM in ("Darwin", "Windows"):
    try:
        converted_count = 0
        for f in os.listdir(OUTPUT_FOLDER):
            if f.lower().endswith('.docx'):
                src = os.path.join(OUTPUT_FOLDER, f)
                pdf_name = os.path.splitext(f)[0] + ".pdf"
                dst = os.path.join(PDF_OUTPUT_FOLDER, pdf_name)
                convert(src, dst)
                if os.path.exists(dst):
                    converted_count += 1
        conversion_ok = converted_count > 0
    except Exception as e:
        print(f"docx2pdf per-file conversion failed: {e}")
        conversion_ok = False

# Fallback: LibreOffice if available (cross-platform: macOS/Windows/Linux)
if not conversion_ok:
    soffice = shutil.which('soffice')
    if soffice:
        try:
            subprocess.run([
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", PDF_OUTPUT_FOLDER,
                *[os.path.join(OUTPUT_FOLDER, f) for f in os.listdir(OUTPUT_FOLDER) if f.lower().endswith('.docx')]
            ], check=True)
            produced = [f for f in os.listdir(PDF_OUTPUT_FOLDER) if f.lower().endswith('.pdf')]
            conversion_ok = len(produced) > 0
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")
            conversion_ok = False

# If conversion worked, delete DOCX folder to keep only PDFs
if conversion_ok:
    try:
        shutil.rmtree(OUTPUT_FOLDER, ignore_errors=True)
    except Exception:
        pass
    print("✅ PDFs created in 'offer_lettersPDF'. 'offer_letters' folder removed.")
else:
    if SYSTEM == "Darwin":
        hint = "On macOS, try running from a local non-iCloud folder or ensure LibreOffice is installed."
    elif SYSTEM == "Windows":
        hint = "On Windows, ensure Microsoft Word or LibreOffice is installed."
    else:
        hint = "On Linux, install LibreOffice (soffice) for conversion."
    print(f"❌ PDF conversion did not complete. {hint}")

print("✅ Process completed. End result: PDFs in 'offer_lettersPDF'.")
